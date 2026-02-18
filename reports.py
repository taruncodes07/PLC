# reports.py

import streamlit as st
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION_START
from docx.shared import RGBColor
import pandas as pd
from fpdf import FPDF
import plotly.express as px
# NOTE: The dashboard module needs to be imported here to use its functions
from dashboard import calculate_kpis, generate_insights, create_filters 
from auth import check_role
import io
from datetime import datetime


def _safe_div(numerator, denominator):
    return (numerator / denominator) if denominator else 0


def _build_text_bars(series, max_width=24):
    if series.empty:
        return []
    max_value = series.max()
    if max_value == 0:
        max_value = 1
    bars = []
    for label, value in series.items():
        bar_len = int((value / max_value) * max_width)
        bars.append((label, value, "#" * bar_len))
    return bars


def _derive_report_metrics(df):
    total_production = df['Actual_Production_Units'].sum()
    total_planned = df['Planned_Production_Units'].sum()
    total_raw_used = df['Raw_Material_Used_kg'].sum()
    total_waste = df['Waste_Weight_kg'].sum()
    total_downtime = df['Downtime_Minutes'].sum()
    total_run_time = df['Total_Time_Run_Minutes'].sum()

    efficiency = _safe_div(total_production, total_planned)
    yield_rate = _safe_div((total_raw_used - total_waste), total_raw_used)
    utilization = _safe_div(total_run_time, (total_run_time + total_downtime))

    date_min = df['Date'].min()
    date_max = df['Date'].max()
    days_covered = df['Date'].dt.date.nunique()

    return {
        "total_production": total_production,
        "total_planned": total_planned,
        "total_raw_used": total_raw_used,
        "total_waste": total_waste,
        "total_downtime": total_downtime,
        "total_run_time": total_run_time,
        "efficiency": efficiency,
        "yield_rate": yield_rate,
        "utilization": utilization,
        "date_min": date_min,
        "date_max": date_max,
        "days_covered": days_covered,
    }


def _build_aggregations(df):
    daily = df.groupby('Date').agg({
        'Actual_Production_Units': 'sum',
        'Planned_Production_Units': 'sum',
        'Downtime_Minutes': 'sum',
        'Waste_Weight_kg': 'sum',
        'Raw_Material_Used_kg': 'sum'
    }).reset_index()
    daily['Efficiency'] = daily.apply(
        lambda row: _safe_div(row['Actual_Production_Units'], row['Planned_Production_Units']), axis=1
    )
    daily['Yield'] = daily.apply(
        lambda row: _safe_div(row['Raw_Material_Used_kg'] - row['Waste_Weight_kg'], row['Raw_Material_Used_kg']), axis=1
    )

    product = df.groupby('Product_Name').agg({
        'Actual_Production_Units': 'sum',
        'Planned_Production_Units': 'sum',
        'Downtime_Minutes': 'sum',
        'Waste_Weight_kg': 'sum',
        'Raw_Material_Used_kg': 'sum'
    }).reset_index()
    product['Efficiency'] = product.apply(
        lambda row: _safe_div(row['Actual_Production_Units'], row['Planned_Production_Units']), axis=1
    )
    product['Waste_Rate'] = product.apply(
        lambda row: _safe_div(row['Waste_Weight_kg'], row['Raw_Material_Used_kg']), axis=1
    )
    product['Share'] = product['Actual_Production_Units'] / product['Actual_Production_Units'].sum()

    shift = df.groupby('Shift').agg({
        'Actual_Production_Units': 'sum',
        'Planned_Production_Units': 'sum',
        'Downtime_Minutes': 'sum',
        'Waste_Weight_kg': 'sum',
        'Raw_Material_Used_kg': 'sum'
    }).reset_index()
    shift['Efficiency'] = shift.apply(
        lambda row: _safe_div(row['Actual_Production_Units'], row['Planned_Production_Units']), axis=1
    )
    shift['Waste_Rate'] = shift.apply(
        lambda row: _safe_div(row['Waste_Weight_kg'], row['Raw_Material_Used_kg']), axis=1
    )
    shift['Downtime_per_Unit'] = shift.apply(
        lambda row: _safe_div(row['Downtime_Minutes'], row['Actual_Production_Units']), axis=1
    )

    operator = df.groupby('Machine_Operator_ID').agg({
        'Actual_Production_Units': 'sum',
        'Planned_Production_Units': 'sum',
        'Downtime_Minutes': 'sum'
    }).reset_index()
    operator['Efficiency'] = operator.apply(
        lambda row: _safe_div(row['Actual_Production_Units'], row['Planned_Production_Units']), axis=1
    )

    downtime = df.groupby('Downtime_Reason')['Downtime_Minutes'].sum().sort_values(ascending=False).reset_index()

    return daily, product, shift, operator, downtime

# --- PDF Generation (using FPDF) ---

class PDF(FPDF):
    """Custom PDF class for report generation."""
    def header(self):
        self.set_fill_color(30, 30, 30) # Dark background feel
        self.set_text_color(255, 255, 255)
        self.rect(0, 0, self.w, 16, 'F')

        self.set_xy(0, 3)
        self.set_font('Arial', 'B', 16)
        self.cell(self.w, 7, 'Production Report', 0, 1, 'C')

        self.set_x(0)
        self.set_text_color(150, 150, 150)
        self.set_font('Arial', '', 10)
        self.cell(self.w, 5, f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}", 0, 1, 'C')

        self.ln(4)
        self.set_text_color(0, 0, 0) # Reset to black for content

    def footer(self):
        self.set_y(-15)
        self.set_font('Arial', 'I', 8)
        self.set_text_color(100, 100, 100)
        self.cell(0, 10, f'Page {self.page_no()}', 0, 0, 'C')

def generate_pdf_report(df_filtered, kpis, insights):
    """Generates the PDF report content."""
    
    pdf = PDF()
    pdf.add_page()

    metrics = _derive_report_metrics(df_filtered)
    daily, product, shift, operator, downtime = _build_aggregations(df_filtered)
    
    # Title Page/Summary
    pdf.set_font('Arial', 'B', 14)
    pdf.cell(0, 10, 'I. Executive Summary', 0, 1, 'L')
    pdf.set_font('Arial', '', 10)
    
    # Insights (Replacing formatting for text-only PDF)
    insight_text = insights.replace(" | ", "\n- ").replace("**", "")
    pdf.multi_cell(0, 5, insight_text + "\n\n", 0, 'L')

    pdf.set_font('Arial', 'B', 11)
    pdf.cell(0, 6, 'Data Coverage', 0, 1, 'L')
    pdf.set_font('Arial', '', 10)
    pdf.multi_cell(
        0,
        5,
        (
            f"Period: {metrics['date_min'].strftime('%Y-%m-%d')} to {metrics['date_max'].strftime('%Y-%m-%d')} | "
            f"Days Covered: {metrics['days_covered']} | Rows: {len(df_filtered)}\n"
            f"Shifts: {df_filtered['Shift'].nunique()} | Products: {df_filtered['Product_Name'].nunique()} | "
            f"Operators: {df_filtered['Machine_Operator_ID'].nunique()}"
        ),
        0,
        'L'
    )

    # KPIs
    pdf.set_font('Arial', 'B', 14)
    pdf.cell(0, 10, 'II. Key Performance Indicators', 0, 1, 'L')
    pdf.set_font('Arial', '', 10)
    
    col_width = pdf.w / 3.5
    for i, (k, v) in enumerate(kpis.items()):
        if i % 2 == 0:
            pdf.ln(5)
            
        pdf.cell(col_width, 6, f"{k}:", 0, 0, 'L')
        pdf.set_font('Arial', 'B', 10)
        pdf.cell(col_width, 6, v, 0, 0, 'L')
        pdf.set_font('Arial', '', 10)

    pdf.ln(6)
    pdf.set_font('Arial', 'B', 11)
    pdf.cell(0, 6, 'Additional KPI Highlights', 0, 1, 'L')
    pdf.set_font('Arial', '', 10)
    avg_daily_prod = _safe_div(metrics['total_production'], metrics['days_covered'])
    avg_daily_downtime = _safe_div(metrics['total_downtime'], metrics['days_covered'])
    pdf.multi_cell(
        0,
        5,
        (
            f"Average Daily Output: {avg_daily_prod:,.0f} units | "
            f"Average Daily Downtime: {avg_daily_downtime:,.1f} min\n"
            f"Plan Attainment: {metrics['efficiency']:.2%} | "
            f"Material Yield: {metrics['yield_rate']:.2%} | "
            f"Utilization: {metrics['utilization']:.2%}"
        ),
        0,
        'L'
    )

    # Tables - Top 5 Downtime Reasons
    pdf.add_page()
    pdf.set_font('Arial', 'B', 14)
    pdf.cell(0, 10, 'III. Top 5 Downtime Reasons', 0, 1, 'L')
    
    df_top_dt = downtime.head(5)
    
    # Table headers
    pdf.set_fill_color(220, 220, 220)
    pdf.set_font('Arial', 'B', 10)
    col_widths = [pdf.w * 0.4, pdf.w * 0.3]
    pdf.cell(col_widths[0], 7, 'Downtime Reason', 1, 0, 'C', 1)
    pdf.cell(col_widths[1], 7, 'Total Minutes', 1, 1, 'C', 1)
    
    # Table rows
    pdf.set_font('Arial', '', 10)
    for index, row in df_top_dt.iterrows():
        pdf.cell(col_widths[0], 6, row['Downtime_Reason'], 1, 0, 'L')
        pdf.cell(col_widths[1], 6, f"{row['Downtime_Minutes']:,.0f}", 1, 1, 'R')

    pdf.ln(6)
    pdf.set_font('Arial', 'B', 12)
    pdf.cell(0, 8, 'IV. Daily Production Trend (Last 10 Days)', 0, 1, 'L')
    pdf.set_font('Arial', '', 9)
    daily_recent = daily.sort_values('Date').tail(10)
    pdf.set_fill_color(230, 230, 230)
    pdf.set_font('Arial', 'B', 9)
    trend_cols = [pdf.w * 0.2, pdf.w * 0.25, pdf.w * 0.2, pdf.w * 0.2]
    pdf.cell(trend_cols[0], 6, 'Date', 1, 0, 'C', 1)
    pdf.cell(trend_cols[1], 6, 'Production', 1, 0, 'C', 1)
    pdf.cell(trend_cols[2], 6, 'Downtime', 1, 0, 'C', 1)
    pdf.cell(trend_cols[3], 6, 'Efficiency', 1, 1, 'C', 1)
    pdf.set_font('Arial', '', 9)
    for _, row in daily_recent.iterrows():
        pdf.cell(trend_cols[0], 6, row['Date'].strftime('%Y-%m-%d'), 1, 0, 'L')
        pdf.cell(trend_cols[1], 6, f"{row['Actual_Production_Units']:,.0f}", 1, 0, 'R')
        pdf.cell(trend_cols[2], 6, f"{row['Downtime_Minutes']:,.0f}", 1, 0, 'R')
        pdf.cell(trend_cols[3], 6, f"{row['Efficiency']:.2%}", 1, 1, 'R')

    pdf.ln(6)
    pdf.set_font('Arial', 'B', 12)
    pdf.cell(0, 8, 'V. Production Peaks (ASCII Visualization)', 0, 1, 'L')
    top_days = daily.sort_values('Actual_Production_Units', ascending=False).head(6)
    for label, value, bar in _build_text_bars(
        top_days.set_index('Date')['Actual_Production_Units']
    ):
        pdf.cell(0, 5, f"{label.strftime('%Y-%m-%d')}: {value:,.0f} | {bar}", 0, 1, 'L')

    pdf.add_page()
    pdf.set_font('Arial', 'B', 14)
    pdf.cell(0, 10, 'VI. Product Mix & Performance', 0, 1, 'L')
    pdf.set_font('Arial', '', 10)
    top_products = product.sort_values('Actual_Production_Units', ascending=False).head(8)
    pdf.set_fill_color(220, 220, 220)
    pdf.set_font('Arial', 'B', 9)
    prod_cols = [pdf.w * 0.3, pdf.w * 0.2, pdf.w * 0.2, pdf.w * 0.15]
    pdf.cell(prod_cols[0], 6, 'Product', 1, 0, 'C', 1)
    pdf.cell(prod_cols[1], 6, 'Units', 1, 0, 'C', 1)
    pdf.cell(prod_cols[2], 6, 'Efficiency', 1, 0, 'C', 1)
    pdf.cell(prod_cols[3], 6, 'Share', 1, 1, 'C', 1)
    pdf.set_font('Arial', '', 9)
    for _, row in top_products.iterrows():
        pdf.cell(prod_cols[0], 6, row['Product_Name'], 1, 0, 'L')
        pdf.cell(prod_cols[1], 6, f"{row['Actual_Production_Units']:,.0f}", 1, 0, 'R')
        pdf.cell(prod_cols[2], 6, f"{row['Efficiency']:.2%}", 1, 0, 'R')
        pdf.cell(prod_cols[3], 6, f"{row['Share']:.1%}", 1, 1, 'R')

    pdf.ln(6)
    pdf.set_font('Arial', 'B', 12)
    pdf.cell(0, 8, 'VII. Shift Performance', 0, 1, 'L')
    pdf.set_font('Arial', 'B', 9)
    shift_cols = [pdf.w * 0.2, pdf.w * 0.2, pdf.w * 0.2, pdf.w * 0.2]
    pdf.set_fill_color(220, 220, 220)
    pdf.cell(shift_cols[0], 6, 'Shift', 1, 0, 'C', 1)
    pdf.cell(shift_cols[1], 6, 'Units', 1, 0, 'C', 1)
    pdf.cell(shift_cols[2], 6, 'Efficiency', 1, 0, 'C', 1)
    pdf.cell(shift_cols[3], 6, 'Downtime/Unit', 1, 1, 'C', 1)
    pdf.set_font('Arial', '', 9)
    for _, row in shift.iterrows():
        pdf.cell(shift_cols[0], 6, str(row['Shift']), 1, 0, 'L')
        pdf.cell(shift_cols[1], 6, f"{row['Actual_Production_Units']:,.0f}", 1, 0, 'R')
        pdf.cell(shift_cols[2], 6, f"{row['Efficiency']:.2%}", 1, 0, 'R')
        pdf.cell(shift_cols[3], 6, f"{row['Downtime_per_Unit']:.3f}", 1, 1, 'R')

    pdf.ln(6)
    pdf.set_font('Arial', 'B', 12)
    pdf.cell(0, 8, 'VIII. Operator Performance (Top 6)', 0, 1, 'L')
    top_ops = operator.sort_values('Actual_Production_Units', ascending=False).head(6)
    pdf.set_fill_color(220, 220, 220)
    pdf.set_font('Arial', 'B', 9)
    op_cols = [pdf.w * 0.3, pdf.w * 0.2, pdf.w * 0.2, pdf.w * 0.2]
    pdf.cell(op_cols[0], 6, 'Operator', 1, 0, 'C', 1)
    pdf.cell(op_cols[1], 6, 'Units', 1, 0, 'C', 1)
    pdf.cell(op_cols[2], 6, 'Efficiency', 1, 0, 'C', 1)
    pdf.cell(op_cols[3], 6, 'Downtime', 1, 1, 'C', 1)
    pdf.set_font('Arial', '', 9)
    for _, row in top_ops.iterrows():
        pdf.cell(op_cols[0], 6, str(row['Machine_Operator_ID']), 1, 0, 'L')
        pdf.cell(op_cols[1], 6, f"{row['Actual_Production_Units']:,.0f}", 1, 0, 'R')
        pdf.cell(op_cols[2], 6, f"{row['Efficiency']:.2%}", 1, 0, 'R')
        pdf.cell(op_cols[3], 6, f"{row['Downtime_Minutes']:,.0f}", 1, 1, 'R')

    pdf.add_page()
    pdf.set_font('Arial', 'B', 14)
    pdf.cell(0, 10, 'IX. Quality & Waste', 0, 1, 'L')
    waste_by_product = product.sort_values('Waste_Rate', ascending=False).head(6)
    pdf.set_font('Arial', 'B', 9)
    pdf.set_fill_color(220, 220, 220)
    waste_cols = [pdf.w * 0.3, pdf.w * 0.2, pdf.w * 0.2, pdf.w * 0.2]
    pdf.cell(waste_cols[0], 6, 'Product', 1, 0, 'C', 1)
    pdf.cell(waste_cols[1], 6, 'Waste (kg)', 1, 0, 'C', 1)
    pdf.cell(waste_cols[2], 6, 'Waste Rate', 1, 0, 'C', 1)
    pdf.cell(waste_cols[3], 6, 'Yield', 1, 1, 'C', 1)
    pdf.set_font('Arial', '', 9)
    for _, row in waste_by_product.iterrows():
        pdf.cell(waste_cols[0], 6, row['Product_Name'], 1, 0, 'L')
        pdf.cell(waste_cols[1], 6, f"{row['Waste_Weight_kg']:,.1f}", 1, 0, 'R')
        pdf.cell(waste_cols[2], 6, f"{row['Waste_Rate']:.2%}", 1, 0, 'R')
        yield_rate = 1 - row['Waste_Rate']
        pdf.cell(waste_cols[3], 6, f"{yield_rate:.2%}", 1, 1, 'R')

    pdf.ln(6)
    pdf.set_font('Arial', 'B', 12)
    pdf.cell(0, 8, 'X. Recommendations & Actions', 0, 1, 'L')
    pdf.set_font('Arial', '', 10)
    recommendations = []
    if metrics['efficiency'] < 0.95:
        recommendations.append("Review planning accuracy and line balancing to improve plan attainment.")
    if metrics['yield_rate'] < 0.97:
        recommendations.append("Investigate material losses and tighten quality control checkpoints.")
    if metrics['total_downtime'] > 0:
        top_reason = downtime.iloc[0]['Downtime_Reason'] if not downtime.empty else "unknown causes"
        recommendations.append(f"Focus downtime reduction on {top_reason} through preventive maintenance and SOP refresh.")
    if shift['Downtime_per_Unit'].max() > shift['Downtime_per_Unit'].mean() * 1.2:
        recommendations.append("Standardize best practices across shifts to reduce variability.")
    if not recommendations:
        recommendations.append("Maintain current operating practices and continue monitoring key drivers.")
    for item in recommendations:
        pdf.multi_cell(0, 5, f"- {item}", 0, 'L')

    pdf.ln(6)
    pdf.set_font('Arial', 'B', 12)
    pdf.cell(0, 8, 'XI. Appendix - Descriptive Statistics', 0, 1, 'L')
    stats = df_filtered[['Actual_Production_Units', 'Planned_Production_Units', 'Downtime_Minutes', 'Waste_Weight_kg']].describe()
    stats = stats.round(2)
    pdf.set_font('Arial', 'B', 8)
    pdf.set_fill_color(220, 220, 220)
    stat_cols = [pdf.w * 0.2] * 5
    headers = ['Metric', 'Mean', 'Std', 'Min', 'Max']
    for i, header in enumerate(headers):
        pdf.cell(stat_cols[i], 6, header, 1, 0, 'C', 1)
    pdf.ln()
    pdf.set_font('Arial', '', 8)
    for metric_name, row in stats.iterrows():
        pdf.cell(stat_cols[0], 6, metric_name, 1, 0, 'L')
        pdf.cell(stat_cols[1], 6, f"{row['mean']:,.2f}", 1, 0, 'R')
        pdf.cell(stat_cols[2], 6, f"{row['std']:,.2f}", 1, 0, 'R')
        pdf.cell(stat_cols[3], 6, f"{row['min']:,.2f}", 1, 0, 'R')
        pdf.cell(stat_cols[4], 6, f"{row['max']:,.2f}", 1, 1, 'R')

    pdf_output = pdf.output()
    if isinstance(pdf_output, (bytes, bytearray)):
        return bytes(pdf_output)
    return pdf_output.encode('latin-1')

# --- DOCX Generation (using python-docx) ---

def generate_docx_report(df_filtered, kpis, insights):
    """Generates the DOCX report content."""
    
    document = Document()

    metrics = _derive_report_metrics(df_filtered)
    daily, product, shift, operator, downtime = _build_aggregations(df_filtered)
    
    # Set Default Font/Size
    style = document.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)
    
    # Title Page
    document.add_heading('Weekly Production Report', 0)
    document.add_paragraph(f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    document.add_paragraph(f"Data Period: {df_filtered['Date'].min().strftime('%Y-%m-%d')} to {df_filtered['Date'].max().strftime('%Y-%m-%d')}")
    document.add_paragraph(
        f"Coverage: {len(df_filtered)} rows | {metrics['days_covered']} days | "
        f"{df_filtered['Shift'].nunique()} shifts | {df_filtered['Product_Name'].nunique()} products | "
        f"{df_filtered['Machine_Operator_ID'].nunique()} operators"
    )
    
    # I. Executive Summary
    document.add_section(WD_SECTION_START.NEW_PAGE)
    document.add_heading('I. Executive Summary', level=1)
    
    # Add Insights as a bulleted list
    p = document.add_paragraph()
    for insight in insights.replace("**", "").split(" | "):
        p.add_run(f'• {insight.strip()}').bold = True
        p = document.add_paragraph()

    document.add_paragraph(
        f"Plan Attainment: {metrics['efficiency']:.2%} | Material Yield: {metrics['yield_rate']:.2%} | "
        f"Utilization: {metrics['utilization']:.2%}"
    )
    
    # II. KPIs
    document.add_heading('II. Key Performance Indicators', level=1)
    table = document.add_table(rows=len(kpis.items()) // 2 + 1, cols=4)
    table.style = 'Medium Shading 1 Accent 1'
    
    # Add KPI data to the table
    kpi_list = list(kpis.items())
    for i in range(len(kpi_list) // 2):
        row = table.rows[i].cells
        
        # KPI 1
        row[0].text = kpi_list[i*2][0]
        row[1].text = kpi_list[i*2][1]
        
        # KPI 2
        row[2].text = kpi_list[i*2 + 1][0]
        row[3].text = kpi_list[i*2 + 1][1]

    # Handle odd number of KPIs
    if len(kpi_list) % 2 != 0:
        row = table.rows[-1].cells
        row[0].text = kpi_list[-1][0]
        row[1].text = kpi_list[-1][1]
        row[2].text = ''
        row[3].text = ''

    # III. Downtime Analysis Table
    document.add_heading('III. Top Downtime Reasons', level=1)
    df_downtime = downtime.head(5)
    
    table_dt = document.add_table(df_downtime.shape[0] + 1, df_downtime.shape[1])
    table_dt.style = 'Light Grid'
    
    # Add header row
    for j, col in enumerate(df_downtime.columns):
        table_dt.cell(0, j).text = col.replace('_', ' ').title()
    
    # Add data rows
    for i, row in df_downtime.iterrows():
        for j, col in enumerate(df_downtime.columns):
            table_dt.cell(i+1, j).text = f"{row[col]:,.0f}" if 'Minutes' in col else str(row[col])

    # IV. Daily Production Trend (Last 10 Days)
    document.add_heading('IV. Daily Production Trend (Last 10 Days)', level=1)
    daily_recent = daily.sort_values('Date').tail(10)
    table_daily = document.add_table(daily_recent.shape[0] + 1, 4)
    table_daily.style = 'Light Grid'
    table_daily.cell(0, 0).text = 'Date'
    table_daily.cell(0, 1).text = 'Production'
    table_daily.cell(0, 2).text = 'Downtime'
    table_daily.cell(0, 3).text = 'Efficiency'
    for i, row in daily_recent.iterrows():
        table_daily.cell(i + 1, 0).text = row['Date'].strftime('%Y-%m-%d')
        table_daily.cell(i + 1, 1).text = f"{row['Actual_Production_Units']:,.0f}"
        table_daily.cell(i + 1, 2).text = f"{row['Downtime_Minutes']:,.0f}"
        table_daily.cell(i + 1, 3).text = f"{row['Efficiency']:.2%}"

    # V. Production Peaks (ASCII Visualization)
    document.add_heading('V. Production Peaks (ASCII Visualization)', level=1)
    top_days = daily.sort_values('Actual_Production_Units', ascending=False).head(6)
    for label, value, bar in _build_text_bars(
        top_days.set_index('Date')['Actual_Production_Units']
    ):
        document.add_paragraph(f"{label.strftime('%Y-%m-%d')}: {value:,.0f} | {bar}")

    # VI. Product Mix & Performance
    document.add_heading('VI. Product Mix & Performance', level=1)
    top_products = product.sort_values('Actual_Production_Units', ascending=False).head(8)
    table_prod = document.add_table(top_products.shape[0] + 1, 4)
    table_prod.style = 'Light Grid'
    table_prod.cell(0, 0).text = 'Product'
    table_prod.cell(0, 1).text = 'Units'
    table_prod.cell(0, 2).text = 'Efficiency'
    table_prod.cell(0, 3).text = 'Share'
    for i, row in top_products.iterrows():
        table_prod.cell(i + 1, 0).text = row['Product_Name']
        table_prod.cell(i + 1, 1).text = f"{row['Actual_Production_Units']:,.0f}"
        table_prod.cell(i + 1, 2).text = f"{row['Efficiency']:.2%}"
        table_prod.cell(i + 1, 3).text = f"{row['Share']:.1%}"

    # VII. Shift Performance
    document.add_heading('VII. Shift Performance', level=1)
    table_shift = document.add_table(shift.shape[0] + 1, 4)
    table_shift.style = 'Light Grid'
    table_shift.cell(0, 0).text = 'Shift'
    table_shift.cell(0, 1).text = 'Units'
    table_shift.cell(0, 2).text = 'Efficiency'
    table_shift.cell(0, 3).text = 'Downtime/Unit'
    for i, row in shift.iterrows():
        table_shift.cell(i + 1, 0).text = str(row['Shift'])
        table_shift.cell(i + 1, 1).text = f"{row['Actual_Production_Units']:,.0f}"
        table_shift.cell(i + 1, 2).text = f"{row['Efficiency']:.2%}"
        table_shift.cell(i + 1, 3).text = f"{row['Downtime_per_Unit']:.3f}"

    # VIII. Operator Performance (Top 6)
    document.add_heading('VIII. Operator Performance (Top 6)', level=1)
    top_ops = operator.sort_values('Actual_Production_Units', ascending=False).head(6)
    table_ops = document.add_table(top_ops.shape[0] + 1, 4)
    table_ops.style = 'Light Grid'
    table_ops.cell(0, 0).text = 'Operator'
    table_ops.cell(0, 1).text = 'Units'
    table_ops.cell(0, 2).text = 'Efficiency'
    table_ops.cell(0, 3).text = 'Downtime'
    for i, row in top_ops.iterrows():
        table_ops.cell(i + 1, 0).text = str(row['Machine_Operator_ID'])
        table_ops.cell(i + 1, 1).text = f"{row['Actual_Production_Units']:,.0f}"
        table_ops.cell(i + 1, 2).text = f"{row['Efficiency']:.2%}"
        table_ops.cell(i + 1, 3).text = f"{row['Downtime_Minutes']:,.0f}"

    # IX. Quality & Waste
    document.add_heading('IX. Quality & Waste', level=1)
    waste_by_product = product.sort_values('Waste_Rate', ascending=False).head(6)
    table_waste = document.add_table(waste_by_product.shape[0] + 1, 4)
    table_waste.style = 'Light Grid'
    table_waste.cell(0, 0).text = 'Product'
    table_waste.cell(0, 1).text = 'Waste (kg)'
    table_waste.cell(0, 2).text = 'Waste Rate'
    table_waste.cell(0, 3).text = 'Yield'
    for i, row in waste_by_product.iterrows():
        table_waste.cell(i + 1, 0).text = row['Product_Name']
        table_waste.cell(i + 1, 1).text = f"{row['Waste_Weight_kg']:,.1f}"
        table_waste.cell(i + 1, 2).text = f"{row['Waste_Rate']:.2%}"
        table_waste.cell(i + 1, 3).text = f"{(1 - row['Waste_Rate']):.2%}"

    # X. Recommendations & Actions
    document.add_heading('X. Recommendations & Actions', level=1)
    recommendations = []
    if metrics['efficiency'] < 0.95:
        recommendations.append("Review planning accuracy and line balancing to improve plan attainment.")
    if metrics['yield_rate'] < 0.97:
        recommendations.append("Investigate material losses and tighten quality control checkpoints.")
    if metrics['total_downtime'] > 0:
        top_reason = downtime.iloc[0]['Downtime_Reason'] if not downtime.empty else "unknown causes"
        recommendations.append(f"Focus downtime reduction on {top_reason} through preventive maintenance and SOP refresh.")
    if shift['Downtime_per_Unit'].max() > shift['Downtime_per_Unit'].mean() * 1.2:
        recommendations.append("Standardize best practices across shifts to reduce variability.")
    if not recommendations:
        recommendations.append("Maintain current operating practices and continue monitoring key drivers.")
    for item in recommendations:
        document.add_paragraph(f"- {item}")

    # XI. Appendix - Descriptive Statistics
    document.add_heading('XI. Appendix - Descriptive Statistics', level=1)
    stats = df_filtered[['Actual_Production_Units', 'Planned_Production_Units', 'Downtime_Minutes', 'Waste_Weight_kg']].describe().round(2)
    table_stats = document.add_table(stats.shape[0] + 1, 5)
    table_stats.style = 'Light Grid'
    table_stats.cell(0, 0).text = 'Metric'
    table_stats.cell(0, 1).text = 'Mean'
    table_stats.cell(0, 2).text = 'Std'
    table_stats.cell(0, 3).text = 'Min'
    table_stats.cell(0, 4).text = 'Max'
    for i, (metric_name, row) in enumerate(stats.iterrows()):
        table_stats.cell(i + 1, 0).text = metric_name
        table_stats.cell(i + 1, 1).text = f"{row['mean']:,.2f}"
        table_stats.cell(i + 1, 2).text = f"{row['std']:,.2f}"
        table_stats.cell(i + 1, 3).text = f"{row['min']:,.2f}"
        table_stats.cell(i + 1, 4).text = f"{row['max']:,.2f}"

    # Save to a BytesIO object
    buffer = io.BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer.read()


def reports_page():
    """Streamlit page for generating and exporting reports."""
    st.title("📄 Export Reports (PDF / Word)")
    st.markdown("---")
    
    if not check_role('Analyst'):
        st.error("Access Denied: You must be an Analyst or Admin to generate reports.")
        return

    if 'df' not in st.session_state or st.session_state['df'].empty:
        st.warning("No dataset loaded. Please load data first.")
        return
        
    st.info("The reports will be generated based on the current filters applied in the Dashboard.")
    
    # Get filtered data, KPIs, and Insights
    # NOTE: Need to suppress filter creation in reports page
    df_filtered = create_filters(st.session_state['df'].copy()) 

    kpis = calculate_kpis(df_filtered)
    insights = generate_insights(df_filtered)

    file_name = st.text_input("Report File Name (e.g., Weekly_Report_2025-12)", "Weekly_Report")
    export_type = st.selectbox("Select Export Type", ["PDF", "Word (.docx)"])
    
    if st.button("Generate & Download Report", type="primary"):
        if not file_name:
            st.error("Please enter a file name.")
            return

        if export_type == "PDF":
            with st.spinner("Generating PDF Report..."):
                pdf_bytes = generate_pdf_report(df_filtered, kpis, insights)
                st.download_button(
                    label="Download PDF",
                    data=pdf_bytes,
                    file_name=f"{file_name}.pdf",
                    mime="application/pdf"
                )
                st.success("PDF generated successfully.")

        elif export_type == "Word (.docx)":
            with st.spinner("Generating Word Report..."):
                docx_bytes = generate_docx_report(df_filtered, kpis, insights)
                st.download_button(
                    label="Download DOCX",
                    data=docx_bytes,
                    file_name=f"{file_name}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
                st.success("Word report generated successfully.")